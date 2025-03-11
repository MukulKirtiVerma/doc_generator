"""
Document Service - Generates formatted documents from Google Vision OCR data
"""
import os
from flask import current_app
import json
from datetime import datetime


def generate_document_from_vision(structured_data, ocr_text, output_format, output_path):
    """
    Generate a document using only Google Vision API data

    Args:
        structured_data (dict): Structured data from Google Vision API
        ocr_text (str): Raw OCR text as fallback if structured data is incomplete
        output_format (str): The desired output format ('docx', 'pdf', 'xlsx')
        output_path (str): The path where the output document should be saved

    Returns:
        tuple: (success, error_message)
    """
    try:
        if output_format == 'docx':
            return generate_word_document(structured_data, ocr_text, output_path)
        elif output_format == 'pdf':
            return generate_pdf_document(structured_data, ocr_text, output_path)
        elif output_format == 'xlsx':
            return generate_excel_document(structured_data, ocr_text, output_path)
        else:
            return False, f"Unsupported output format: {output_format}"
    except Exception as e:
        return False, f"Error generating document: {str(e)}"


def generate_word_document(structured_data, ocr_text, output_path):
    """Generate Word document from structured Vision API data"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # If structured data is missing or empty, use raw OCR text as fallback
        if not structured_data or not structured_data.get('pages'):
            paragraph = doc.add_paragraph(ocr_text)
            doc.save(output_path)
            return True, None

        # Process each page
        for page in structured_data.get('pages', []):
            # Sort blocks by vertical position (top to bottom)
            blocks = page.get('blocks', [])
            if not blocks:
                continue

            # Sort blocks by vertical position for proper reading order
            sorted_blocks = sorted(blocks, key=lambda b: min(v[1] for v in b.get('bounding_box', [(0, 0)])))

            for block in sorted_blocks:
                if block.get('type') == 'table':
                    # Create a table in Word
                    table_data = next((t for t in structured_data.get('tables', [])
                                       if t.get('page') == structured_data.get('pages').index(page) and
                                       t.get('block') == blocks.index(block)), None)

                    if table_data and table_data.get('rows'):
                        rows = table_data.get('rows', [])
                        if rows:
                            # Determine the max columns in any row
                            max_cols = max(len(row) for row in rows)
                            table = doc.add_table(rows=len(rows), cols=max_cols)
                            table.style = 'Table Grid'

                            for i, row in enumerate(rows):
                                for j, cell_text in enumerate(row):
                                    if j < len(table.rows[i].cells):
                                        table.rows[i].cells[j].text = cell_text

                            # Add spacing after table
                            doc.add_paragraph()
                else:
                    # Process text block
                    for paragraph in block.get('paragraphs', []):
                        p = doc.add_paragraph()

                        # Get text
                        text = paragraph.get('text', '')
                        if not text.strip():
                            continue

                        # Detect if paragraph is a heading (simplified heuristic)
                        if len(text) < 100 and (text.strip().endswith(':') or text.isupper() or
                                                any(text.startswith(h) for h in ["Chapter ", "Section "])):
                            p.style = 'Heading 2'
                            p.add_run(text)
                        else:
                            # Add text with potential formatting
                            for word in paragraph.get('words', []):
                                word_text = word.get('text', '')
                                run = p.add_run(word_text + ' ')

                                # Apply basic formatting if detected
                                # This is very simple detection based on common patterns
                                if word_text.isupper() and len(word_text) > 1:
                                    run.bold = True

                                # Check for likely emphasis patterns
                                if word_text.startswith('*') and word_text.endswith('*'):
                                    run.italic = True

                                # Check for likely underline patterns
                                if word_text.startswith('_') and word_text.endswith('_'):
                                    run.underline = True

        # Save the document
        doc.save(output_path)
        return True, None

    except ImportError:
        # Handle missing python-docx library
        return False, "python-docx library is required but not installed. Install it with: pip install python-docx"
    except Exception as e:
        return False, f"Error generating Word document: {str(e)}"


# Update the generate_excel_document function to accept the ocr_text parameter
def generate_excel_document(structured_data, ocr_text, output_path):
    """Generate Excel document from structured Vision API data"""
    try:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active

        # Find tables in the document
        tables = structured_data.get('tables', [])

        if tables:
            # Use the first table for the main worksheet
            main_table = tables[0]

            for row_idx, row in enumerate(main_table.get('rows', []), 1):
                for col_idx, cell_text in enumerate(row, 1):
                    ws.cell(row=row_idx, column=col_idx, value=cell_text)

            # Create additional worksheets for other tables
            for table_idx, table in enumerate(tables[1:], 1):
                ws_name = f"Table {table_idx + 1}"
                ws = wb.create_sheet(title=ws_name)

                for row_idx, row in enumerate(table.get('rows', []), 1):
                    for col_idx, cell_text in enumerate(row, 1):
                        ws.cell(row=row_idx, column=col_idx, value=cell_text)
        else:
            # No tables found, try to create a structured spreadsheet from blocks
            row_idx = 1

            # If structured data is missing or empty, use raw OCR text as fallback
            if not structured_data or not structured_data.get('pages'):
                ws.cell(row=row_idx, column=1, value=ocr_text)
            else:
                for page in structured_data.get('pages', []):
                    for block in page.get('blocks', []):
                        for paragraph in block.get('paragraphs', []):
                            text = paragraph.get('text', '')
                            if text.strip():
                                ws.cell(row=row_idx, column=1, value=text)
                                row_idx += 1

        # Save the workbook
        wb.save(output_path)
        return True, None
    except Exception as e:
        return False, f"Error generating Excel document: {str(e)}"


# Update the generate_pdf_document function to accept the ocr_text parameter
def generate_pdf_document(structured_data, ocr_text, output_path):
    """Generate PDF document from structured Vision API data"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
        from reportlab.lib.styles import getSampleStyleSheet

        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []

        # If structured data is missing or empty, use raw OCR text as fallback
        if not structured_data or not structured_data.get('pages'):
            elements.append(Paragraph(ocr_text, styles['Normal']))
            doc.build(elements)
            return True, None

        # Process each page
        for page in structured_data.get('pages', []):
            # Sort blocks by vertical position
            blocks = page.get('blocks', [])
            blocks.sort(key=lambda b: min(v[1] for v in b.get('bounding_box', [(0, 0)])))

            for block in blocks:
                if block.get('type') == 'table':
                    # Create a table in PDF
                    table_data = next((t for t in structured_data.get('tables', [])
                                       if t.get('page') == structured_data.get('pages').index(page) and
                                       t.get('block') == blocks.index(block)),
                                      None)

                    if table_data and table_data.get('rows'):
                        t = Table(table_data.get('rows'))
                        elements.append(t)
                        elements.append(Spacer(1, 12))
                else:
                    # Process text block
                    for paragraph in block.get('paragraphs', []):
                        text = paragraph.get('text', '')
                        if text.strip():
                            # Apply appropriate style
                            if len(text) < 100 and text.strip().endswith(':'):
                                p = Paragraph(text, styles['Heading2'])
                            else:
                                p = Paragraph(text, styles['Normal'])

                            elements.append(p)
                            elements.append(Spacer(1, 6))

        # Build the PDF
        doc.build(elements)
        return True, None
    except Exception as e:
        return False, f"Error generating PDF document: {str(e)}"